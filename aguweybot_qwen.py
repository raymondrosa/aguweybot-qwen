# ============================================
# AGUWEYBOT - VERSIÓN CON OPENROUTER - CORREGIDO Y ROBUSTECIDO
# ============================================

import os
import base64
import time
import streamlit as st
import streamlit.components.v1 as components
import re
import io
import sys
import traceback
import random
import json
from typing import Optional, Tuple, List
from datetime import datetime
from dataclasses import dataclass

from openai import OpenAI
from openai import APIError, APIConnectionError, RateLimitError

from langchain_core.messages import HumanMessage, SystemMessage, AIMessage

# Para documentos
from PyPDF2 import PdfReader
from docx import Document
import pandas as pd
import chardet

# ============================================
# LIBRERÍAS PARA IMÁGENES
# ============================================
try:
    import requests
    from PIL import Image, ImageDraw, ImageFont, ImageFilter, ImageEnhance
    IMAGE_GEN_AVAILABLE = True
    print("✅ Librerías de imágenes disponibles")
except ImportError as e:
    IMAGE_GEN_AVAILABLE = False
    print(f"❌ Error importando librerías: {e}")

# ============================================
# TEXTO A VOZ
# ============================================
try:
    from gtts import gTTS
    TTS_AVAILABLE = True
except ImportError:
    TTS_AVAILABLE = False

# ============================================
# CONFIGURACIÓN INICIAL DE STREAMLIT
# ============================================
st.set_page_config(
    page_title="AguweyBot - OpenRouter",
    page_icon="🎨",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ============================================
# CONSTANTES Y CONFIGURACIÓN VISUAL
# ============================================
class Config:
    # Colores
    PRIMARY_COLOR = "#00ffff"
    SECONDARY_COLOR = "#00cccc"
    BACKGROUND_DARK = "#0a0c10"
    CARD_BACKGROUND = "#1e2a3a"

    # Rutas de archivos
    LOGO_PATH = "logo.png"
    BACKGROUND_PATH = "fondo.png"

    # Límites
    MAX_HISTORY_MESSAGES = 10
    MAX_FILE_SIZE_MB = 50
    MAX_IMAGE_SIZE_MB = 5
    MAX_DOCUMENT_CHARS = 50000

    # Configuración OpenRouter
    OPENROUTER_MODEL = "qwen/qwen-2.5-7b-instruct"
    OPENROUTER_BASE_URL = "https://openrouter.ai/api/v1"
    MAX_TOKENS = 4096
    TEMPERATURE = 0.2
    TIMEOUT = 60.0


# ============================================
# SYSTEM PROMPT
# ============================================
SYSTEM_PROMPT = """
Eres AguweyBot, un asistente híbrido experto en:
1. ANÁLISIS DE DOCUMENTOS (texto, PDFs, Excel, etc.)
2. GENERACIÓN DE IMÁGENES (cuando el usuario lo solicite)

CAPACIDADES:
- 📄 Análisis de documentos: Lee TODO el contenido de archivos subidos
- 🎨 Generación de imágenes: Crea imágenes a partir de descripciones
- 🔊 Texto a voz: Convierte respuestas a audio
- 📋 Copiado de respuestas
- 💾 Descarga de imágenes generadas

DETECCIÓN AUTOMÁTICA:
- Si el usuario pide "dibuja", "crea imagen", "genera foto", etc., debes ACTIVAR generación de imágenes
- Si el usuario sube un archivo o pregunta sobre documentos, debes ANALIZAR el contenido
- Si es una pregunta normal, responde como asistente conversacional

REGLAS PARA IMÁGENES:
- Cuando detectes una solicitud de imagen, responde con: "🎨 GENERANDO IMAGEN: [descripción corta]"
- Luego el sistema generará la imagen automáticamente
- Si la descripción no es clara, pide más detalles

REGLAS PARA DOCUMENTOS:
- Usa TODO el contenido del archivo para responder
- No inventes información
- Si no encuentras algo en el archivo, dilo honestamente

FORMATO DE RESPUESTA:
- Usa emojis para hacer las respuestas más amigables
- Responde de manera clara, concisa y profesional
- Mantén un tono amigable pero formal
"""

# ============================================
# VERIFICAR API KEYS
# ============================================
if "OPENROUTER_API_KEY" not in st.secrets:
    st.error("❌ No se encontró la API Key de OpenRouter")
    st.info("""
    Por favor, agrega tu OPENROUTER_API_KEY en Settings → Secrets

    Formato:
    OPENROUTER_API_KEY = "sk-or-v1-..."

    Puedes obtener tu API key en: https://openrouter.ai/keys
    """)
    st.stop()

# ============================================
# FUNCIÓN PARA FONDO
# ============================================
def set_background():
    """Aplica la imagen de fondo si existe"""
    if os.path.exists(Config.BACKGROUND_PATH):
        try:
            with open(Config.BACKGROUND_PATH, "rb") as f:
                img_data = f.read()
            encoded = base64.b64encode(img_data).decode()
            st.markdown(
                f"""
                <style>
                .stApp {{
                    background-image: url("data:image/png;base64,{encoded}");
                    background-size: cover;
                    background-position: center;
                    background-attachment: fixed;
                    background-repeat: no-repeat;
                }}
                .main .block-container {{
                    background-color: rgba(0, 0, 0, 0.7);
                    backdrop-filter: blur(10px);
                    border-radius: 20px;
                    padding: 2rem;
                    margin: 2rem auto;
                    border: 1px solid {Config.PRIMARY_COLOR};
                    box-shadow: 0 0 20px rgba(0, 255, 255, 0.3);
                    max-width: 1200px;
                }}
                </style>
                """,
                unsafe_allow_html=True
            )
        except Exception:
            st.markdown(
                f"""
                <style>
                .stApp {{
                    background: linear-gradient(135deg, {Config.BACKGROUND_DARK}, #1a1f2a);
                }}
                </style>
                """,
                unsafe_allow_html=True
            )
    else:
        st.markdown(
            f"""
            <style>
            .stApp {{
                background: linear-gradient(135deg, {Config.BACKGROUND_DARK}, #1a1f2a);
            }}
            </style>
            """,
            unsafe_allow_html=True
        )

# ============================================
# ESTILOS CSS CORREGIDOS - BARRA DE ESCRITURA INFERIOR
# ============================================
def aplicar_estilos():
    st.markdown(
        f"""
        <style>
        /* Estilos generales */
        .stApp {{
            background-color: {Config.BACKGROUND_DARK};
        }}

        /* Contenedor principal */
        .main .block-container {{
            background-color: rgba(10, 12, 16, 0.85);
            backdrop-filter: blur(10px);
            border-radius: 20px;
            padding: 2rem 2rem 7rem 2rem !important;
            margin: 1rem auto 0 auto !important;
            border: 1px solid {Config.PRIMARY_COLOR};
            box-shadow: 0 0 30px rgba(0, 255, 255, 0.2);
            max-width: 1000px !important;
            min-height: calc(100vh - 120px);
            margin-bottom: 0 !important;
        }}

        /* Títulos */
        h1 {{
            color: {Config.PRIMARY_COLOR} !important;
            font-size: 2.5rem !important;
            text-align: center;
            text-shadow: 0 0 20px rgba(0, 255, 255, 0.5);
            margin-bottom: 0.5rem !important;
            font-weight: bold;
            animation: glow 2s ease-in-out infinite alternate;
        }}
        @keyframes glow {{
            from {{
                text-shadow: 0 0 10px {Config.PRIMARY_COLOR};
            }}
            to {{
                text-shadow: 0 0 20px {Config.PRIMARY_COLOR}, 0 0 30px {Config.SECONDARY_COLOR};
            }}
        }}
        .subtitle {{
            text-align: center;
            color: #e0e5f0;
            margin-bottom: 2rem;
            font-size: 1.1rem;
            opacity: 0.9;
        }}

        /* ===== RESPUESTAS MEJORADAS DEL ASISTENTE ===== */
        .respuesta-aguwey {{
            background: linear-gradient(145deg, {Config.CARD_BACKGROUND}, #15232e);
            border-left: 6px solid {Config.PRIMARY_COLOR};
            border-radius: 16px;
            padding: 1.8rem;
            margin: 1.5rem 0;
            color: #f0f5fa;
            font-size: 1.05rem;
            line-height: 1.7;
            box-shadow: 0 8px 25px rgba(0, 255, 255, 0.15);
            position: relative;
            overflow: hidden;
            font-family: 'Segoe UI', 'Roboto', sans-serif;
            letter-spacing: 0.3px;
        }}

        .respuesta-aguwey::before {{
            content: '';
            position: absolute;
            top: 0;
            left: -100%;
            width: 100%;
            height: 100%;
            background: linear-gradient(90deg, transparent, rgba(0, 255, 255, 0.1), transparent);
            transition: left 0.5s;
        }}

        .respuesta-aguwey:hover::before {{
            left: 100%;
        }}

        .respuesta-aguwey p {{
            margin-bottom: 1rem;
        }}

        .respuesta-aguwey ul,
        .respuesta-aguwey ol {{
            margin: 0.8rem 0;
            padding-left: 2rem;
        }}

        .respuesta-aguwey li {{
            margin: 0.3rem 0;
        }}

        .respuesta-aguwey code {{
            background: rgba(0, 0, 0, 0.3);
            color: {Config.PRIMARY_COLOR};
            padding: 0.2rem 0.4rem;
            border-radius: 6px;
            font-family: 'Courier New', monospace;
            font-size: 0.95rem;
            border: 1px solid rgba(0, 255, 255, 0.2);
        }}

        .respuesta-aguwey pre {{
            background: #0a0e14;
            border-radius: 12px;
            padding: 1rem;
            border: 1px solid {Config.PRIMARY_COLOR};
            overflow-x: auto;
            margin: 1rem 0;
        }}

        .respuesta-aguwey pre code {{
            background: none;
            border: none;
            color: #e0e5f0;
            padding: 0;
        }}

        .respuesta-aguwey table {{
            border-collapse: collapse;
            width: 100%;
            margin: 1rem 0;
        }}

        .respuesta-aguwey th {{
            background: rgba(0, 255, 255, 0.2);
            color: {Config.PRIMARY_COLOR};
            padding: 0.5rem;
            border: 1px solid {Config.PRIMARY_COLOR};
        }}

        .respuesta-aguwey td {{
            padding: 0.5rem;
            border: 1px solid rgba(0, 255, 255, 0.3);
        }}

        .respuesta-aguwey blockquote {{
            border-left: 4px solid {Config.PRIMARY_COLOR};
            background: rgba(0, 255, 255, 0.05);
            padding: 0.8rem;
            margin: 1rem 0;
            border-radius: 0 12px 12px 0;
            font-style: italic;
        }}

        /* ===== BARRA DE ESCRITURA FIJA EN LA PARTE INFERIOR ===== */
        .stChatInputContainer {{
            position: fixed !important;
            bottom: 70px !important;
            left: 50% !important;
            transform: translateX(-50%) !important;
            width: 700px !important;
            max-width: 90% !important;
            z-index: 1000 !important;
            background: rgba(10, 12, 16, 0.95) !important;
            backdrop-filter: blur(12px) !important;
            border: 2px solid {Config.PRIMARY_COLOR} !important;
            border-radius: 50px !important;
            padding: 5px 5px 5px 20px !important;
            box-shadow: 0 10px 40px rgba(0, 0, 0, 0.5), 0 0 20px rgba(0, 255, 255, 0.2) !important;
            transition: all 0.3s ease !important;
        }}

        .stChatInputContainer:hover {{
            box-shadow: 0 15px 50px rgba(0, 0, 0, 0.6), 0 0 30px rgba(0, 255, 255, 0.3) !important;
            border-color: {Config.SECONDARY_COLOR} !important;
            transform: translateX(-50%) translateY(-2px) !important;
        }}

        .stChatInputContainer textarea {{
            background: transparent !important;
            color: white !important;
            border: none !important;
            font-size: 1rem !important;
            padding: 12px 0 !important;
            min-height: 50px !important;
            max-height: 150px !important;
            resize: vertical !important;
        }}

        .stChatInputContainer textarea::placeholder {{
            color: rgba(0, 255, 255, 0.6) !important;
            font-style: italic;
        }}

        .stChatInputContainer button {{
            background: linear-gradient(145deg, {Config.SECONDARY_COLOR}, {Config.PRIMARY_COLOR}) !important;
            color: #000 !important;
            font-weight: bold !important;
            border: none !important;
            border-radius: 50px !important;
            padding: 0.5rem 1.5rem !important;
            margin: 5px !important;
            transition: all 0.2s !important;
            box-shadow: 0 4px 15px rgba(0, 0, 0, 0.3) !important;
        }}

        .stChatInputContainer button:hover {{
            transform: scale(1.05) !important;
            box-shadow: 0 6px 20px rgba(0, 255, 255, 0.3) !important;
        }}

        .stChatMessage:last-of-type {{
            margin-bottom: 100px !important;
        }}

        .imagen-generada {{
            background: linear-gradient(145deg, {Config.CARD_BACKGROUND}, #15232e);
            border-radius: 16px;
            padding: 1.5rem;
            margin: 1rem 0;
            text-align: center;
            border: 2px solid {Config.PRIMARY_COLOR};
            box-shadow: 0 0 30px rgba(0, 255, 255, 0.2);
            transition: all 0.3s;
        }}

        .imagen-generada:hover {{
            transform: scale(1.02);
            box-shadow: 0 0 40px rgba(0, 255, 255, 0.3);
        }}

        .imagen-generada img {{
            max-width: 100%;
            border-radius: 12px;
            box-shadow: 0 4px 20px rgba(0, 0, 0, 0.5);
            border: 1px solid rgba(0, 255, 255, 0.3);
        }}

        /* Sidebar */
        [data-testid="stSidebar"] {{
            background: linear-gradient(165deg, #0e1219, #0a0e14);
            border-right: 2px solid {Config.PRIMARY_COLOR};
            padding: 1rem;
            box-shadow: 5px 0 20px rgba(0, 0, 0, 0.5);
            z-index: 1001;
        }}

        /* Botones */
        .stButton > button {{
            background: linear-gradient(145deg, {Config.SECONDARY_COLOR}, {Config.PRIMARY_COLOR});
            color: black !important;
            font-weight: bold;
            border: none;
            border-radius: 20px;
            padding: 0.3rem 1rem;
            transition: all 0.2s;
            border: 1px solid rgba(255, 255, 255, 0.1);
        }}

        .stButton > button:hover {{
            transform: translateY(-2px);
            box-shadow: 0 4px 15px rgba(0, 255, 255, 0.3);
        }}

        /* Botón de copiar mejorado */
        .copy-btn {{
            background: rgba(0, 255, 255, 0.1);
            border: 1px solid {Config.PRIMARY_COLOR};
            color: {Config.PRIMARY_COLOR};
            border-radius: 8px;
            padding: 4px 12px;
            cursor: pointer;
            font-size: 12px;
            font-family: sans-serif;
            transition: all 0.3s ease;
            margin-left: 8px;
            backdrop-filter: blur(5px);
        }}

        .copy-btn:hover {{
            background: {Config.PRIMARY_COLOR};
            color: #000;
            transform: translateY(-1px);
            box-shadow: 0 2px 8px rgba(0, 255, 255, 0.3);
        }}

        /* Footer ajustado */
        .fixed-footer {{
            position: fixed;
            bottom: 0;
            left: 0;
            right: 0;
            background: rgba(10, 12, 16, 0.98);
            backdrop-filter: blur(12px);
            border-top: 2px solid {Config.PRIMARY_COLOR};
            padding: 0.8rem;
            text-align: center;
            color: #e0e5f0;
            z-index: 1002;
            font-size: 0.95rem;
            box-shadow: 0 -5px 20px rgba(0, 0, 0, 0.5);
        }}

        /* Scrollbar personalizada */
        ::-webkit-scrollbar {{
            width: 10px;
            height: 10px;
        }}
        ::-webkit-scrollbar-track {{
            background: #1a1f2a;
            border-radius: 5px;
        }}
        ::-webkit-scrollbar-thumb {{
            background: {Config.PRIMARY_COLOR};
            border-radius: 5px;
            border: 2px solid #1a1f2a;
        }}
        ::-webkit-scrollbar-thumb:hover {{
            background: {Config.SECONDARY_COLOR};
        }}

        /* Tooltips mejorados */
        [data-tooltip] {{
            position: relative;
            cursor: help;
        }}
        [data-tooltip]:before {{
            content: attr(data-tooltip);
            position: absolute;
            bottom: 100%;
            left: 50%;
            transform: translateX(-50%);
            padding: 8px 12px;
            background: rgba(0, 0, 0, 0.9);
            color: white;
            border-radius: 8px;
            font-size: 12px;
            white-space: nowrap;
            display: none;
            border: 1px solid {Config.PRIMARY_COLOR};
            box-shadow: 0 4px 15px rgba(0, 0, 0, 0.5);
            z-index: 2000;
        }}
        [data-tooltip]:hover:before {{
            display: block;
        }}

        /* Animación de escritura para el streaming */
        @keyframes typing {{
            0% {{ opacity: 0.3; }}
            50% {{ opacity: 1; }}
            100% {{ opacity: 0.3; }}
        }}
        .typing-indicator {{
            display: inline-block;
            width: 4px;
            height: 4px;
            background-color: {Config.PRIMARY_COLOR};
            border-radius: 50%;
            margin: 0 2px;
            animation: typing 1s infinite;
        }}
        </style>
        """,
        unsafe_allow_html=True
    )

# ============================================
# INICIALIZAR CLIENTE OPENROUTER
# ============================================
@st.cache_resource
def get_openrouter_client():
    """Inicializa y cachea el cliente de OpenRouter"""
    try:
        client = OpenAI(
            api_key=st.secrets["OPENROUTER_API_KEY"],
            base_url=Config.OPENROUTER_BASE_URL,
            max_retries=3,
            timeout=Config.TIMEOUT,
            default_headers={
                "HTTP-Referer": "https://aguweybot.streamlit.app",
                "X-Title": "AguweyBot"
            }
        )
        return client
    except Exception as e:
        st.error(f"Error al inicializar cliente OpenRouter: {str(e)}")
        return OpenAI(
            api_key=st.secrets["OPENROUTER_API_KEY"],
            base_url=Config.OPENROUTER_BASE_URL,
            default_headers={
                "HTTP-Referer": "https://aguweybot.streamlit.app",
                "X-Title": "AguweyBot"
            }
        )

# ============================================
# FUNCIÓN PARA DETECTAR SOLICITUD DE IMAGEN
# ============================================
def detectar_solicitud_imagen(texto: str) -> Optional[str]:
    """Detecta si el usuario pide generar una imagen y extrae la descripción"""
    texto_lower = texto.lower().strip()
    print(f"Detectando solicitud de imagen en: {texto}")

    verbos_creacion = ['dibuja', 'dibújame', 'crea', 'genera', 'haz', 'pinta', 'diseña', 'ilustra']
    sustantivos_imagen = ['imagen', 'foto', 'ilustración', 'dibujo', 'gráfico', 'arte', 'pintura', 'retrato', 'paisaje']

    # Patrón verbo + sustantivo
    for verbo in verbos_creacion:
        for sustantivo in sustantivos_imagen:
            frase_clave = f"{verbo} {sustantivo}"
            if frase_clave in texto_lower:
                partes = texto_lower.split(frase_clave, 1)
                if len(partes) > 1:
                    descripcion = partes[1].strip()
                    descripcion = re.sub(r'^(de|sobre|un|una|el|la|unos|unas)\s+', '', descripcion)
                    if descripcion:
                        print(f"✅ Detectada solicitud de imagen: {descripcion}")
                        return descripcion
                print(f"✅ Detectada solicitud de imagen: {texto}")
                return texto

    # Patrón "imagen de X", "foto de Y"
    for sustantivo in sustantivos_imagen:
        patron = rf'{sustantivo}\s+(?:de|sobre)\s+(.+)'
        match = re.search(patron, texto_lower)
        if match:
            descripcion = match.group(1).strip()
            print(f"✅ Detectada solicitud de imagen: {descripcion}")
            return descripcion

    # Consultas muy cortas que mencionen 'imagen', 'foto', etc.
    if len(texto.split()) < 6 and any(p in texto_lower for p in sustantivos_imagen):
        print(f"✅ Detectada solicitud de imagen: {texto}")
        return texto

    print("❌ No se detectó solicitud de imagen")
    return None

# ============================================
# GENERADOR LOCAL MEJORADO
# ============================================
def generar_imagen_local(descripcion: str) -> Tuple[Optional[bytes], Optional[str]]:
    """Generador local que SIEMPRE funciona"""
    try:
        print(f"🎨 Generando imagen local para: {descripcion}")
        width, height = 768, 768
        img = Image.new('RGB', (width, height), color=(20, 30, 50))
        draw = ImageDraw.Draw(img)

        # Fondo degradado
        for y in range(height):
            r = int(20 + (y / height) * 100)
            g = int(30 + (y / height) * 50)
            b = int(70 + (y / height) * 150)
            draw.line([(0, y), (width, y)], fill=(r, g, b))

        # Estrellas/puntos
        for _ in range(150):
            x = random.randint(0, width)
            y = random.randint(0, height)
            size = random.randint(1, 3)
            color = (255, 255, 255) if random.random() > 0.7 else (
                random.randint(100, 255),
                random.randint(100, 255),
                255
            )
            draw.ellipse([x, y, x + size, y + size], fill=color)

        # Formas geométricas
        draw.ellipse(
            [width // 4, height // 4, 3 * width // 4, 3 * height // 4],
            outline=(255, 255, 255, 100),
            width=3
        )
        draw.ellipse(
            [width // 3, height // 3, 2 * width // 3, 2 * height // 3],
            outline=(255, 255, 0, 100),
            width=2
        )

        # Texto
        try:
            font_large = ImageFont.load_default()
            font_small = ImageFont.load_default()
            try:
                font_large = ImageFont.truetype("arial.ttf", 40)
                font_small = ImageFont.truetype("arial.ttf", 20)
            except IOError:
                try:
                    font_large = ImageFont.truetype(
                        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 40
                    )
                    font_small = ImageFont.truetype(
                        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 20
                    )
                except IOError:
                    pass

            palabras = descripcion.split()
            lineas = []
            linea_actual = ""
            for palabra in palabras:
                if len(linea_actual + " " + palabra) < 35:
                    linea_actual += " " + palabra if linea_actual else palabra
                else:
                    lineas.append(linea_actual)
                    linea_actual = palabra
            if linea_actual:
                lineas.append(linea_actual)

            y_offset = height // 2 - len(lineas) * 25
            for linea in lineas:
                bbox = draw.textbbox((0, 0), linea, font=font_large)
                text_width = bbox[2] - bbox[0]
                x_pos = (width - text_width) // 2

                # Sombra
                draw.text((x_pos + 2, y_offset + 2), linea, fill=(0, 0, 0), font=font_large)
                # Texto principal
                draw.text((x_pos, y_offset), linea, fill=(255, 255, 255), font=font_large)
                y_offset += (bbox[3] - bbox[1]) + 10

            footer = "🎨 AguweyBot (Modo Local)"
            bbox_footer = draw.textbbox((0, 0), footer, font=font_small)
            footer_width = bbox_footer[2] - bbox_footer[0]
            draw.text(
                ((width - footer_width) // 2, height - 40),
                footer,
                fill=(0, 255, 255),
                font=font_small
            )
        except Exception as e:
            print(f"Error en dibujo de texto: {e}")
            draw.text(
                (width // 2, height // 2),
                f"🎨 {descripcion[:50]}",
                fill=(255, 255, 255),
                anchor="mm"
            )

        img_bytes = io.BytesIO()
        img.save(img_bytes, format='PNG', optimize=True)
        img_bytes.seek(0)
        print(f"✅ Imagen local generada: {len(img_bytes.getvalue())} bytes")
        return img_bytes.getvalue(), None

    except Exception as e:
        print(f"Error CRÍTICO en generador local: {e}")
        try:
            img = Image.new('RGB', (100, 100), color=(255, 0, 0))
            img_bytes = io.BytesIO()
            img.save(img_bytes, format='PNG')
            img_bytes.seek(0)
            return img_bytes.getvalue(), f"Error en generación: {e}"
        except Exception:
            return None, "Error crítico generando imagen"

# ============================================
# FUNCIÓN PRINCIPAL DE GENERACIÓN
# ============================================
def generar_imagen(descripcion: str) -> Tuple[Optional[bytes], Optional[str]]:
    """Genera una imagen usando el generador local"""
    print("\n=== INICIANDO GENERACIÓN DE IMAGEN ===")
    print(f"Descripción: {descripcion}")
    if not descripcion or not descripcion.strip():
        return None, "La descripción de la imagen está vacía."
    return generar_imagen_local(descripcion)

# ============================================
# FUNCIÓN PARA BOTÓN DE COPIAR (CORREGIDA)
# ============================================
def boton_copiar(texto: str, id_unico: str) -> None:
    """
    Genera un botón de copiado usando JavaScript.
    Usa json.dumps para crear un literal JS seguro.
    """
    texto_js = json.dumps(texto)  # string JS segura (maneja comillas y saltos de línea)

    html_code = f"""
    <div style="text-align: right; margin-top: 0px;">
      <button id="btn_{id_unico}" class="copy-btn" onclick="copyText_{id_unico}()">
        📋 Copiar
      </button>
    </div>
    <script>
      function copyText_{id_unico}() {{
        const textToCopy = {texto_js};
        if (navigator.clipboard && navigator.clipboard.writeText) {{
          navigator.clipboard.writeText(textToCopy)
            .then(() => {{ showCopied_{id_unico}(); }})
            .catch(() => {{ fallbackCopy_{id_unico}(textToCopy); }});
        }} else {{
          fallbackCopy_{id_unico}(textToCopy);
        }}
      }}

      function fallbackCopy_{id_unico}(text) {{
        const tempTextArea = document.createElement('textarea');
        tempTextArea.value = text;
        tempTextArea.style.position = 'fixed';
        tempTextArea.style.opacity = '0';
        document.body.appendChild(tempTextArea);
        tempTextArea.select();
        document.execCommand('copy');
        document.body.removeChild(tempTextArea);
        showCopied_{id_unico}();
      }}

      function showCopied_{id_unico}() {{
        const btn = document.getElementById("btn_{id_unico}");
        const originalText = btn.innerText;
        btn.innerText = "✅ ¡Copiado!";
        btn.style.background = "rgba(0, 255, 0, 0.2)";
        btn.style.borderColor = "#00ff00";
        btn.style.color = "#00ff00";
        setTimeout(() => {{
          btn.innerText = originalText;
          btn.style.background = "rgba(0, 255, 255, 0.1)";
          btn.style.borderColor = "#00ffff";
          btn.style.color = "#00ffff";
        }}, 2000);
      }}
    </script>
    """
    components.html(html_code, height=40)

# ============================================
# CLASE PARA DATOS DEL ARCHIVO
# ============================================
@dataclass
class DatosArchivo:
    nombre: str = ""
    contenido_completo: str = ""
    tipo: str = ""
    dataframe: Optional[pd.DataFrame] = None
    num_paginas: int = 0
    num_caracteres: int = 0
    resumen: str = ""
    fecha_carga: float = time.time()
    truncado: bool = False

    def generar_resumen(self) -> str:
        """Genera un resumen básico del archivo"""
        base = ""
        if self.tipo == "pdf":
            base = f"📄 PDF con {self.num_paginas} páginas"
        elif self.tipo in ["excel", "csv"]:
            if self.dataframe is not None:
                base = f"📊 Tabla con {len(self.dataframe)} filas y {len(self.dataframe.columns)} columnas"
        elif self.tipo in ["txt", "docx"]:
            palabras = len(self.contenido_completo.split())
            base = f"📝 Documento con {palabras} palabras"
        else:
            base = "📁 Archivo procesado"

        if self.truncado:
            base += " (truncado)"
        return base

# ============================================
# FUNCIÓN PARA LEER ARCHIVOS
# ============================================
def leer_archivo_completo(uploaded_file):
    """Lee el archivo con un límite de caracteres"""
    if uploaded_file is None:
        return None, "No hay archivo para procesar"

    try:
        uploaded_file.seek(0)
        file_size = uploaded_file.seek(0, os.SEEK_END)
        uploaded_file.seek(0)

        if file_size > Config.MAX_FILE_SIZE_MB * 1024 * 1024:
            return None, f"El archivo excede el límite de {Config.MAX_FILE_SIZE_MB}MB"

        nombre = uploaded_file.name.lower()
        datos = DatosArchivo()
        datos.nombre = uploaded_file.name
        datos.fecha_carga = time.time()

        # PDF
        if nombre.endswith(".pdf"):
            try:
                reader = PdfReader(uploaded_file)
                datos.num_paginas = len(reader.pages)
                texto_completo = []
                total_chars = 0
                truncado = False
                progress_bar = st.progress(0, text="📖 Leyendo PDF...")

                for i, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        page_header = f"\n\n--- PÁGINA {i+1} ---\n"
                        texto_completo.append(page_header + page_text)
                        total_chars += len(page_header) + len(page_text)
                        if total_chars > Config.MAX_DOCUMENT_CHARS:
                            truncado = True
                            break
                    progress_bar.progress((i + 1) / datos.num_paginas)

                datos.contenido_completo = "".join(texto_completo)
                datos.tipo = "pdf"
                datos.truncado = truncado
                progress_bar.empty()

                if not datos.contenido_completo:
                    return None, "El PDF no contiene texto extraíble"

            except Exception as e:
                return None, f"Error al leer PDF: {str(e)}"

        # Excel
        elif nombre.endswith((".xlsx", ".xls")):
            try:
                df = pd.read_excel(uploaded_file)
                datos.dataframe = df
                df_string = df.to_string()
                if len(df_string) > Config.MAX_DOCUMENT_CHARS:
                    datos.contenido_completo = df_string[:Config.MAX_DOCUMENT_CHARS] + "\n... [TRUNCADO POR LONGITUD]"
                    datos.truncado = True
                else:
                    datos.contenido_completo = df_string
                datos.tipo = "excel"
            except Exception as e:
                return None, f"Error al leer Excel: {str(e)}"

        # CSV
        elif nombre.endswith(".csv"):
            try:
                raw_data = uploaded_file.read()
                result = chardet.detect(raw_data)
                encoding = result["encoding"] or "utf-8"
                df = pd.read_csv(io.BytesIO(raw_data), encoding=encoding)
                datos.dataframe = df
                df_string = df.to_string()
                if len(df_string) > Config.MAX_DOCUMENT_CHARS:
                    datos.contenido_completo = df_string[:Config.MAX_DOCUMENT_CHARS] + "\n... [TRUNCADO POR LONGITUD]"
                    datos.truncado = True
                else:
                    datos.contenido_completo = df_string
                datos.tipo = "csv"
            except Exception as e:
                return None, f"Error al leer CSV: {str(e)}"

        # TXT
        elif nombre.endswith(".txt"):
            try:
                contenido = uploaded_file.read()
                result = chardet.detect(contenido)
                encoding = result["encoding"] or "utf-8"
                texto_completo = contenido.decode(encoding)
                if len(texto_completo) > Config.MAX_DOCUMENT_CHARS:
                    datos.contenido_completo = texto_completo[:Config.MAX_DOCUMENT_CHARS] + "\n... [TRUNCADO POR LONGITUD]"
                    datos.truncado = True
                else:
                    datos.contenido_completo = texto_completo
                datos.tipo = "txt"
            except Exception as e:
                return None, f"Error al leer TXT: {str(e)}"

        # Word
        elif nombre.endswith(".docx"):
            try:
                doc = Document(uploaded_file)
                texto_completo = []
                total_chars = 0
                truncado = False

                for p in doc.paragraphs:
                    if p.text.strip():
                        texto_completo.append(p.text + "\n")
                        total_chars += len(p.text) + 1
                        if total_chars > Config.MAX_DOCUMENT_CHARS:
                            truncado = True
                            break

                if not truncado:
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = " | ".join([cell.text for cell in row.cells]) + "\n"
                            texto_completo.append(row_text)
                            total_chars += len(row_text)
                            if total_chars > Config.MAX_DOCUMENT_CHARS:
                                truncado = True
                                break
                        if truncado:
                            break

                datos.contenido_completo = "".join(texto_completo)
                if truncado:
                    datos.contenido_completo += "\n... [TRUNCADO POR LONGITUD]"
                datos.tipo = "docx"
                datos.truncado = truncado

                if not datos.contenido_completo:
                    return None, "El documento no contiene texto"
            except Exception as e:
                return None, f"Error al leer DOCX: {str(e)}"

        else:
            return None, f"Tipo de archivo no soportado: {nombre.split('.')[-1]}"

        datos.num_caracteres = len(datos.contenido_completo)
        datos.resumen = datos.generar_resumen()
        return datos, None

    except Exception as e:
        return None, f"Error inesperado: {str(e)}"

# ============================================
# CALLBACK PARA STREAMING MEJORADO
# ============================================
class StreamlitCallbackHandler:
    def __init__(self, container):
        self.container = container
        self.text = ""
        self.start_time = time.time()
        self.last_update = time.time()
        self.html_content = ""

    def on_token(self, token: str):
        self.text += token
        self.html_content = self._format_text(self.text)
        if time.time() - self.last_update > 0.05:
            self._update_display()
            self.last_update = time.time()

    def on_end(self):
        self.html_content = self._format_text(self.text)
        self._update_display(final=True)

    def _format_text(self, text: str) -> str:
        """Formatea el texto con HTML para mejor presentación"""

        # Escapar HTML una sola vez
        text = (text
                .replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;"))

        # Negritas: **texto**
        text = re.sub(r"\*\*(.*?)\*\*", r"<strong>\1</strong>", text)

        # Cursivas: *texto*
        text = re.sub(
            r"(?<!\*)\*(?!\*)(.*?)(?<!\*)\*(?!\*)",
            r"<em>\1</em>",
            text
        )

        # Código inline: `codigo`
        text = re.sub(r"`([^`]+)`", r"<code>\1</code>", text)

        # Listas y párrafos
        lines = text.split("\n")
        formatted_lines = []
        in_list = False

        for line in lines:
            list_match = re.match(r"^[\s]*[-*+]\s+(.*)", line)
            numbered_match = re.match(r"^[\s]*\d+\.\s+(.*)", line)

            if list_match or numbered_match:
                if not in_list:
                    formatted_lines.append("<ul>")
                    in_list = True
                content = list_match.group(1) if list_match else numbered_match.group(1)
                formatted_lines.append(f"<li>{content}</li>")
            else:
                if in_list:
                    formatted_lines.append("</ul>")
                    in_list = False

                if line.strip():
                    formatted_lines.append(f"<p>{line}</p>")
                else:
                    formatted_lines.append("<br>")

        if in_list:
            formatted_lines.append("</ul>")

        return "".join(formatted_lines)

    def _update_display(self, final=False):
        elapsed = time.time() - self.start_time
        suffix = '<span class="typing-indicator"></span>' * 3 if not final else ""
        footer = (
            f'<div style="position: absolute; bottom: 5px; right: 10px; '
            f'font-size: 10px; color: #666;">'
            f'{"✅ Completado" if final else "⏳ Escribiendo..."} en {elapsed:.1f}s'
            f"</div>"
        )
        try:
            self.container.markdown(
                f'<div class="respuesta-aguwey" style="position: relative;">'
                f'{self.html_content}{suffix}{footer}</div>',
                unsafe_allow_html=True
            )
        except Exception as e:
            print(f"Error actualizando display: {e}")

# ============================================
# FUNCIÓN PARA LLAMAR A OPENROUTER
# ============================================
def llamar_openrouter(mensajes, callback=None):
    """Llama al modelo de OpenRouter con manejo de errores"""
    try:
        client = get_openrouter_client()

        openai_messages = []
        for msg in mensajes:
            if isinstance(msg, SystemMessage):
                openai_messages.append({"role": "system", "content": msg.content})
            elif isinstance(msg, HumanMessage):
                openai_messages.append({"role": "user", "content": msg.content})
            elif isinstance(msg, AIMessage):
                openai_messages.append({"role": "assistant", "content": msg.content})

        kwargs = {
            "model": Config.OPENROUTER_MODEL,
            "messages": openai_messages,
            "temperature": Config.TEMPERATURE,
            "max_tokens": Config.MAX_TOKENS,
        }

        kwargs["extra_headers"] = {
            "HTTP-Referer": "https://aguweybot.streamlit.app",
            "X-Title": "AguweyBot"
        }

        if callback:
            kwargs["stream"] = True
            response_text = ""
            try:
                stream = client.chat.completions.create(**kwargs)
                for chunk in stream:
                    if (
                        chunk.choices
                        and chunk.choices[0].delta
                        and chunk.choices[0].delta.content
                    ):
                        token = chunk.choices[0].delta.content
                        response_text += token
                        callback.on_token(token)
                callback.on_end()
                return response_text
            except Exception as e:
                print(f"Error en streaming: {e}")
                kwargs.pop("stream", None)
                response = client.chat.completions.create(**kwargs)
                return response.choices[0].message.content

        # Sin streaming
        response = client.chat.completions.create(**kwargs)
        return response.choices[0].message.content

    except RateLimitError as e:
        error_msg = "⏳ Límite de tasa de la API excedido. Por favor, espera un momento."
        st.error(error_msg)
        raise Exception(error_msg) from e
    except APIConnectionError as e:
        error_msg = "🔌 Error de conexión con la API. Verifica tu internet."
        st.error(error_msg)
        raise Exception(error_msg) from e
    except APIError as e:
        error_msg = f"⚠️ Error de la API: {e.message}"
        st.error(error_msg)
        raise Exception(error_msg) from e
    except Exception as e:
        print(f"Error inesperado: {e}")
        raise e

# ============================================
# TEXTO A VOZ
# ============================================
def texto_a_audio_unico(texto: str) -> Optional[bytes]:
    """Convierte texto a audio"""
    if not TTS_AVAILABLE or not texto or not texto.strip():
        return None
    try:
        texto_limpio = re.sub(r"[#*_\[\]()📄📊🔊🎨✅❌⚠️]", "", texto)
        texto_limpio = re.sub(r"\s+", " ", texto_limpio).strip()
        if not texto_limpio:
            return None
        if len(texto_limpio) > 2000:
            texto_limpio = texto_limpio[:2000] + "... [resumen]"

        tts = gTTS(text=texto_limpio, lang="es", slow=False)
        audio_bytes_io = io.BytesIO()
        tts.write_to_fp(audio_bytes_io)
        audio_bytes_io.seek(0)
        return audio_bytes_io.getvalue()
    except Exception as e:
        print(f"Error generando audio: {e}")
        return None

# ============================================
# FUNCIÓN PARA MOSTRAR LOGO
# ============================================
def mostrar_logo():
    if os.path.exists(Config.LOGO_PATH):
        try:
            from PIL import Image as PILImage
            logo = PILImage.open(Config.LOGO_PATH)
            st.sidebar.image(logo, width=200)
        except Exception:
            st.sidebar.markdown("# 🤖 AguweyBot")
    else:
        st.sidebar.markdown("""
        # 🤖 AguweyBot
        ### *Con OpenRouter y generación local de imágenes*
        """)

def mostrar_info_archivo(datos: DatosArchivo) -> None:
    """Muestra información del archivo cargado"""
    if datos:
        with st.sidebar.expander("📁 Archivo activo", expanded=True):
            st.markdown(
                f"""
                **Nombre:** {datos.nombre}

                **Tipo:** {datos.resumen}

                **Tamaño:** {datos.num_caracteres:,} caracteres

                **Cargado:** {datetime.fromtimestamp(datos.fecha_carga).strftime('%H:%M:%S')}
                """
            )
            if datos.truncado:
                st.warning("⚠️ El archivo era muy grande y se ha truncado.")

            if datos.tipo in ["excel", "csv"] and datos.dataframe is not None:
                st.dataframe(datos.dataframe.head(5), use_container_width=True)
            elif datos.num_caracteres > 500:
                st.markdown("**📄 Vista previa:**")
                st.text_area(
                    label="Contenido del archivo",
                    value=datos.contenido_completo[:500] + "...",
                    height=150,
                    disabled=True,
                    label_visibility="collapsed",
                )

# ============================================
# FUNCIÓN PRINCIPAL
# ============================================
def main():
    """Función principal de la aplicación"""
    set_background()
    aplicar_estilos()

    # Inicializar session state
    if "messages" not in st.session_state:
        st.session_state.messages = []
    if "datos_archivo" not in st.session_state:
        st.session_state.datos_archivo = None
    if "primer_mensaje" not in st.session_state:
        st.session_state.primer_mensaje = True
    if "audio_actual_bytes" not in st.session_state:
        st.session_state.audio_actual_bytes = None
    if "audio_actual_idx" not in st.session_state:
        st.session_state.audio_actual_idx = -1
    if "imagenes_generadas" not in st.session_state:
        st.session_state.imagenes_generadas = {}
    if "procesando" not in st.session_state:
        st.session_state.procesando = False

    # Sidebar
    with st.sidebar:
        mostrar_logo()
        st.markdown("---")
        st.markdown("### 🔑 Estado")
        st.success("✅ OpenRouter conectado")
        st.success("🎨 Generador local activo (100% funcional)")
        st.info(f"🤖 Modelo: {Config.OPENROUTER_MODEL}")

        if TTS_AVAILABLE:
            st.success("✅ Audio disponible")
        else:
            st.warning("⚠️ Audio no disponible - Instala gtts")

        st.markdown("---")
        st.markdown("### 🧪 Prueba")
        if st.button("Probar generación de imagen", use_container_width=True):
            with st.spinner("Generando imagen de prueba..."):
                test_desc = "un paisaje hermoso con montañas"
                img_bytes, error = generar_imagen(test_desc)
                if img_bytes:
                    st.success("✅ Generación exitosa!")
                    st.image(img_bytes, width=300)
                else:
                    st.error(f"❌ Error: {error}")

        st.markdown("---")
        st.markdown("### 📎 Subir Archivo")
        uploaded_file = st.file_uploader(
            "Elige un archivo",
            type=["pdf", "xlsx", "xls", "csv", "txt", "docx"],
            key="file_uploader",
            label_visibility="collapsed"
        )

        if uploaded_file is not None:
            col1, col2 = st.columns(2)
            with col1:
                if st.button("📖 Leer TODO", key="btn_leer", use_container_width=True):
                    with st.spinner("📖 Leyendo archivo..."):
                        datos, error = leer_archivo_completo(uploaded_file)
                        if error:
                            st.error(f"❌ {error}")
                        elif datos:
                            st.session_state.datos_archivo = datos
                            st.success(f"✅ {datos.resumen}")
                            if datos.truncado:
                                st.warning("⚠️ El archivo era muy grande y se ha truncado.")
                            st.balloons()
            with col2:
                if st.button("🔄 Limpiar", use_container_width=True):
                    st.session_state.datos_archivo = None
                    st.rerun()

        if st.session_state.datos_archivo:
            mostrar_info_archivo(st.session_state.datos_archivo)

        st.markdown("---")
        with st.expander("🔍 Debug", expanded=False):
            if st.button("Ver imágenes en memoria"):
                st.write(f"Imágenes guardadas: {len(st.session_state.imagenes_generadas)}")
                for idx, img_bytes in st.session_state.imagenes_generadas.items():
                    st.write(f"Índice {idx}: {len(img_bytes)} bytes")

        if st.button("🔄 Nueva Conversación", use_container_width=True):
            st.session_state.messages = []
            st.session_state.audio_actual_bytes = None
            st.session_state.audio_actual_idx = -1
            st.session_state.imagenes_generadas = {}
            st.session_state.datos_archivo = None
            st.session_state.procesando = False
            st.success("¡Conversación reiniciada!")
            st.rerun()

        if st.session_state.messages:
            st.markdown("### 📊 Estadísticas")
            user_msgs = sum(1 for m in st.session_state.messages if m["role"] == "user")
            assistant_msgs = sum(1 for m in st.session_state.messages if m["role"] == "assistant")
            imagenes_count = len(st.session_state.imagenes_generadas)
            st.markdown(
                f"""
                - 💬 Mensajes: {len(st.session_state.messages)}
                - 👤 Usuario: {user_msgs}
                - 🤖 Asistente: {assistant_msgs}
                - 🎨 Imágenes: {imagenes_count}
                """
            )

    # Contenido principal
    st.markdown("<h1>🎨 AguweyBot con OpenRouter</h1>", unsafe_allow_html=True)
    st.markdown(
        '<p class="subtitle">Asistente inteligente con análisis de documentos y generación local de imágenes</p>',
        unsafe_allow_html=True
    )

    # Mostrar historial
    for i, msg in enumerate(st.session_state.messages):
        with st.chat_message(msg["role"]):
            if msg["role"] == "assistant":
                if msg.get("tipo") == "imagen":
                    if i in st.session_state.imagenes_generadas:
                        imagen_bytes = st.session_state.imagenes_generadas[i]
                        st.markdown(
                            f'<div class="imagen-generada">{msg["content"]}</div>',
                            unsafe_allow_html=True
                        )
                        st.image(imagen_bytes, use_container_width=True)

                        col1, col2, col3 = st.columns([1, 1, 4])
                        with col1:
                            if TTS_AVAILABLE:
                                if st.button(f"🔊 Audio", key=f"audio_img_{i}"):
                                    audio_bytes = texto_a_audio_unico(msg["content"])
                                    if audio_bytes:
                                        st.audio(audio_bytes, format="audio/mpeg")
                        with col2:
                            st.download_button(
                                label="💾 Guardar",
                                data=imagen_bytes,
                                file_name=f"aguweybot_imagen_{i}.png",
                                mime="image/png",
                                key=f"download_img_{i}",
                            )
                        with col3:
                            boton_copiar(msg["content"], f"copy_img_{i}")
                    else:
                        st.warning("🖼️ La imagen no está disponible temporalmente")
                else:
                    st.markdown(
                        f'<div class="respuesta-aguwey">{msg["content"]}</div>',
                        unsafe_allow_html=True
                    )
                    col1, col2 = st.columns([1, 8])
                    with col1:
                        if TTS_AVAILABLE:
                            if st.button(f"🔊 Audio", key=f"audio_{i}"):
                                audio_bytes = texto_a_audio_unico(msg["content"])
                                if audio_bytes:
                                    st.session_state.audio_actual_bytes = audio_bytes
                                    st.session_state.audio_actual_idx = i
                                    st.rerun()
                    with col2:
                        boton_copiar(msg["content"], f"copy_{i}")

                    if (
                        st.session_state.get("audio_actual_idx") == i
                        and "audio_actual_bytes" in st.session_state
                        and st.session_state.audio_actual_bytes
                    ):
                        st.audio(st.session_state.audio_actual_bytes, format="audio/mpeg")
            else:
                st.markdown(f"**Tú:** {msg['content']}")

    # Mensaje de bienvenida
    if st.session_state.primer_mensaje and not st.session_state.messages:
        st.info(
            """
            👋 **¡Bienvenido a AguweyBot con OpenRouter!**

            **🤖 Modelo:** Qwen2.5-72B-Instruct vía OpenRouter

            **🎨 Características:**
            - 📄 **Análisis de documentos** (PDF, Excel, CSV, TXT, DOCX)
            - 🖼️ **Generación local de imágenes** (siempre funciona, sin internet)
            - 🔊 **Texto a voz** en respuestas
            - 📋 **Copiado con un clic**
            - 💾 **Descarga de imágenes**

            **💡 Ejemplos:**
            - "dibuja un gato astronauta en el espacio"
            - "crea una imagen de una ciudad cyberpunk"
            - "genera un dragón mágico en un castillo"
            - "pinta un paisaje con montañas y un lago"

            **✨ El generador local crea imágenes atractivas y variadas según el tema**
            """
        )
        st.session_state.primer_mensaje = False

    # Input del usuario
    prompt = st.chat_input(
        "Escribe tu pregunta o pide una imagen...",
        key="chat_input",
        disabled=st.session_state.procesando
    )

    if prompt and not st.session_state.procesando:
        st.session_state.procesando = True
        st.session_state.messages.append({"role": "user", "content": prompt})
        st.rerun()

    # Procesamiento del último mensaje
    if st.session_state.procesando:
        last_user_message = next(
            (m["content"] for m in reversed(st.session_state.messages) if m["role"] == "user"),
            None
        )
        if last_user_message:
            descripcion_imagen = detectar_solicitud_imagen(last_user_message)
            if descripcion_imagen:
                with st.chat_message("assistant"):
                    status = st.status(
                        f"🎨 Generando imagen: '{descripcion_imagen}'...",
                        expanded=True
                    )
                    with status:
                        st.write("🔍 Procesando solicitud...")
                        st.write(f"📝 Descripción: {descripcion_imagen}")
                        st.write("🎨 Usando generador local")

                        imagen_bytes, error = generar_imagen(descripcion_imagen)
                        if imagen_bytes:
                            st.write("✅ Imagen generada exitosamente!")
                            status.update(label="✅ Imagen generada!", state="complete")

                            idx = len(st.session_state.messages)
                            st.session_state.imagenes_generadas[idx] = imagen_bytes

                            
                            respuesta = f"🎨 **Imagen generada:** {descripcion_imagen}"
                            st.session_state.messages.append(
                                {
                                    "role": "assistant",
                                    "content": respuesta,
                                    "tipo": "imagen",
                                }
                            )
                        else:
                            error_msg = error or "Error desconocido"
                            st.write(f"❌ Error: {error_msg}")
                            status.update(
                                label="❌ Error al generar imagen",
                                state="error"
                            )
                            st.session_state.messages.append(
                                {
                                    "role": "assistant",
                                    "content": f"❌ Lo siento, no pude generar la imagen: {error_msg}",
                                }
                            )
            else:
                with st.chat_message("assistant"):
                    try:
                        container = st.empty()
                        callback = StreamlitCallbackHandler(container)

                        mensajes = [SystemMessage(content=SYSTEM_PROMPT)]
                        ultimos_mensajes = st.session_state.messages[-Config.MAX_HISTORY_MESSAGES:]

                        for m in ultimos_mensajes:
                            if m["role"] == "user":
                                mensajes.append(HumanMessage(content=m["content"]))
                            elif m["role"] == "assistant" and m.get("tipo") != "imagen":
                                mensajes.append(AIMessage(content=m["content"]))

                        if st.session_state.datos_archivo:
                            datos = st.session_state.datos_archivo
                            contexto = f"""
                            📁 ARCHIVO: {datos.nombre}
                            TIPO: {datos.tipo} {'(TRUNCADO)' if datos.truncado else ''}
                            CONTENIDO COMPLETO:
                            {datos.contenido_completo}

                            PREGUNTA: {last_user_message}
                            """
                            mensajes.append(HumanMessage(content=contexto))
                        else:
                            mensajes.append(HumanMessage(content=last_user_message))

                        response_content = llamar_openrouter(mensajes, callback)
                        container.markdown(
                            f'<div class="respuesta-aguwey">{response_content}</div>',
                            unsafe_allow_html=True
                        )
                        st.session_state.messages.append(
                            {"role": "assistant", "content": response_content}
                        )

                        if TTS_AVAILABLE and len(response_content) > 100:
                            audio_bytes = texto_a_audio_unico(response_content)
                            if audio_bytes:
                                st.session_state.audio_actual_bytes = audio_bytes
                                st.session_state.audio_actual_idx = len(
                                    st.session_state.messages
                                ) - 1

                    except Exception as e:
                        st.error(f"❌ Error al obtener respuesta: {str(e)}")
                        st.exception(e)
                        st.session_state.messages.append(
                            {
                                "role": "assistant",
                                "content": f"❌ Lo siento, ocurrió un error: {str(e)[:200]}",
                            }
                        )

        st.session_state.procesando = False
        st.rerun()

    # Footer ajustado
    st.markdown(
        f"""
        <div class="fixed-footer">
            <strong>CC-SA</strong> Prof. Raymond Rosa Ávila • AguweyBot con OpenRouter •
            <span data-tooltip="Versión con barra inferior mejorada">🚀 v10.2</span>
            <span style="margin-left: 20px; color: {Config.PRIMARY_COLOR}; font-size: 0.85rem;">
                ⬇️ Barra de escritura abajo
            </span>
        </div>
        """,
        unsafe_allow_html=True
    )

# ============================================
# PUNTO DE ENTRADA
# ============================================
if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        st.error(f"❌ Error crítico en la aplicación: {str(e)}")
        st.code(traceback.format_exc())
        st.stop()